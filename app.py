#!/usr/bin/env python3
"""
NAKED RESUME - Universal Resume Analyzer
Fixed for OpenAI Project API Keys (sk-proj-)
"""

import os
import re
from typing import Optional, Tuple, Dict
from dataclasses import dataclass
from pathlib import Path

import gradio as gr
from openai import OpenAI, AuthenticationError, APIError
from PyPDF2 import PdfReader
from docx import Document

# =============================================================================
# CONFIGURATION
# =============================================================================

@dataclass
class Config:
    """Application configuration"""
    # YOUR PROJECT API KEY
    OPENAI_API_KEY: str = "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"
    
    # OPTIONAL: If you have organization ID, add it here
    # Get from: https://platform.openai.com/settings/organization/general
    OPENAI_ORG_ID: Optional[str] = None
    
    # OPTIONAL: Project ID (sometimes needed for project keys)
    # Get from your project settings in OpenAI dashboard
    OPENAI_PROJECT_ID: Optional[str] = None
    
    MODEL: str = "gpt-3.5-turbo"  # Using 3.5-turbo for better compatibility
    MAX_TOKENS_ANALYSIS: int = 1500
    MAX_TOKENS_REWRITE: int = 2500
    TEMPERATURE_ANALYSIS: float = 0.2
    TEMPERATURE_REWRITE: float = 0.6

    def is_configured(self) -> bool:
        return self.OPENAI_API_KEY.startswith("sk-") and len(self.OPENAI_API_KEY) > 20


# Initialize configuration
config = Config()

# =============================================================================
# CLIENT INITIALIZATION - FIXED FOR PROJECT KEYS
# =============================================================================

class OpenAIClient:
    """Manages OpenAI client with project key support"""
    
    def __init__(self):
        self.client: Optional[OpenAI] = None
        self.error: Optional[str] = None
        self._init_client()
    
    def _init_client(self):
        if not config.is_configured():
            self.error = "API key not configured"
            print("ERROR: API key not configured")
            return
        
        try:
            # Initialize with project key support
            client_kwargs = {"api_key": config.OPENAI_API_KEY}
            
            # Add organization if available
            if config.OPENAI_ORG_ID:
                client_kwargs["organization"] = config.OPENAI_ORG_ID
            
            # For project keys, we need to set default headers
            if config.OPENAI_PROJECT_ID:
                client_kwargs["default_headers"] = {
                    "OpenAI-Project": config.OPENAI_PROJECT_ID
                }
            
            self.client = OpenAI(**client_kwargs)
            
            # Test with a simple models list call
            try:
                models = self.client.models.list()
                model_ids = [m.id for m in models.data]
                print(f"SUCCESS: OpenAI connected. Available models: {len(model_ids)}")
                print(f"Models include: {model_ids[:3]}...")
            except Exception as e:
                print(f"WARNING: Connected but couldn't list models: {e}")
                # Still mark as ready if we got here
                pass
            
            print("SUCCESS: OpenAI client initialized")
            
        except AuthenticationError as e:
            self.error = f"Authentication failed: {str(e)}"
            print(f"ERROR: Authentication failed - {str(e)}")
            self.client = None
        except Exception as e:
            self.error = f"Connection error: {str(e)}"
            print(f"ERROR: {str(e)}")
            self.client = None
    
    def is_ready(self) -> bool:
        return self.client is not None
    
    def get_client(self) -> Optional[OpenAI]:
        return self.client


# Global client
ai_client = OpenAIClient()

# =============================================================================
# DOCUMENT PROCESSING
# =============================================================================

class DocumentParser:
    """Extract text from PDF and DOCX files"""
    
    @staticmethod
    def extract(file_obj) -> Tuple[Optional[str], Optional[str]]:
        if file_obj is None:
            return None, "No file uploaded"
        
        if hasattr(file_obj, 'name'):
            file_path = file_obj.name
        elif isinstance(file_obj, str):
            file_path = file_obj
        else:
            return None, "Invalid file format"
        
        ext = Path(file_path).suffix.lower()
        
        try:
            if ext == '.pdf':
                return DocumentParser._parse_pdf(file_path)
            elif ext == '.docx':
                return DocumentParser._parse_docx(file_path)
            else:
                return None, f"Unsupported format: {ext}. Use PDF or DOCX."
        except Exception as e:
            return None, f"Parse error: {str(e)}"
    
    @staticmethod
    def _parse_pdf(path: str) -> Tuple[str, None]:
        reader = PdfReader(path)
        texts = []
        for i, page in enumerate(reader.pages, 1):
            text = page.extract_text()
            if text:
                texts.append(f"[Page {i}]\n{text}")
        
        full = "\n\n".join(texts).strip()
        if not full:
            raise ValueError("PDF is empty or image-based")
        return full, None
    
    @staticmethod
    def _parse_docx(path: str) -> Tuple[str, None]:
        doc = Document(path)
        texts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                texts.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    texts.append(row_text)
        
        full = "\n".join(texts).strip()
        if not full:
            raise ValueError("Document is empty")
        return full, None


# =============================================================================
# AI SERVICE
# =============================================================================

class ResumeOptimizer:
    """Core AI logic for analysis and rewriting"""
    
    def __init__(self, client: OpenAIClient):
        self.client = client
    
    def _truncate(self, text: str, limit: int) -> str:
        return text[:limit] + "..." if len(text) > limit else text
    
    def analyze(self, resume: str, job: str) -> Tuple[str, Dict]:
        """Generate match analysis"""
        if not self.client.is_ready():
            return self._error_msg("OpenAI not connected. Check API key."), {}
        
        resume_text = self._truncate(resume, 3000)
        job_text = self._truncate(job, 2000)
        
        system_msg = "You are an expert ATS analyzer and career strategist."
        
        prompt = f"""Analyze this resume against the job description.

MATCH PERCENTAGE: [X]%

TECHNICAL SKILLS MATCH:
- In Resume: [list]
- In Job Only: [missing]
- Match Rate: X%

EXPERIENCE ALIGNMENT:
- Years Required vs Found
- Alignment: [Good/Fair/Poor]

KEY STRENGTHS:
• [Strength 1]
• [Strength 2]

CRITICAL GAPS:
• [Gap 1]
• [Gap 2]

IMMEDIATE ACTIONS:
1. [Action]
2. [Action]
3. [Action]

RESUME:
{resume_text}

JOB:
{job_text}"""

        try:
            resp = self.client.get_client().chat.completions.create(
                model=config.MODEL,
                messages=[
                    {"role": "system", "content": system_msg},
                    {"role": "user", "content": prompt}
                ],
                temperature=config.TEMPERATURE_ANALYSIS,
                max_tokens=config.MAX_TOKENS_ANALYSIS
            )
            content = resp.choices[0].message.content
            
            pct_match = re.search(r'MATCH PERCENTAGE:\s*(\d+)%', content)
            percentage = int(pct_match.group(1)) if pct_match else 0
            
            return content, {"percentage": percentage}
            
        except Exception as e:
            return self._error_msg(f"API Error: {str(e)}"), {}
    
    def rewrite(self, resume: str, job: str) -> str:
        """Rewrite resume with quantified achievements"""
        if not self.client.is_ready():
            return self._error_msg("OpenAI not connected")
        
        resume_text = self._truncate(resume, 3000)
        job_text = self._truncate(job, 2000)
        
        prompt = f"""Rewrite this resume for the job. Requirements:
1. QUANTIFY: Use %, $, numbers (e.g., "Increased revenue by 35%")
2. TONE: Experienced professional, action verbs
3. STRUCTURE: Summary, Skills, Experience, Education

RESUME:
{resume_text}

JOB:
{job_text}

OUTPUT ONLY THE REWRITTEN RESUME:"""

        try:
            resp = self.client.get_client().chat.completions.create(
                model=config.MODEL,
                messages=[{"role": "user", "content": prompt}],
                temperature=config.TEMPERATURE_REWRITE,
                max_tokens=config.MAX_TOKENS_REWRITE
            )
            return resp.choices[0].message.content
            
        except Exception as e:
            return self._error_msg(f"API Error: {str(e)}")
    
    def _error_msg(self, msg: str) -> str:
        return f"## ⚠️ {msg}"


# Initialize optimizer
optimizer = ResumeOptimizer(ai_client)

# =============================================================================
# MAIN APPLICATION
# =============================================================================

class ResumeApp:
    """Main application controller"""
    
    def __init__(self):
        self.parser = DocumentParser()
    
    def process(self, resume_file, job_desc: str) -> Tuple[str, str, str]:
        """Process resume"""
        
        if not ai_client.is_ready():
            error_detail = ai_client.error or "Unknown error"
            return (
                f"## ⚠️ OpenAI Connection Failed\n\n**Error:** {error_detail}\n\n**Solutions:**\n1. Check if your API key is valid at https://platform.openai.com/api-keys\n2. Try creating a new key\n3. Ensure you have billing set up\n4. Check if you need to add OPENAI_ORG_ID or OPENAI_PROJECT_ID",
                "Connection failed",
                "Status: API Error"
            )
        
        if resume_file is None:
            return ("## Upload Required", "Please upload resume", "Status: Waiting")
        
        if not job_desc or len(job_desc.strip()) < 100:
            return ("## Job Description Required", "Paste job description (100+ chars)", "Status: Too short")
        
        resume_text, error = self.parser.extract(resume_file)
        if error:
            return (f"## Parse Error\n\n{error}", "File error", "Status: Failed")
        
        if len(resume_text) < 200:
            return ("## Content Error", "Resume too short", "Status: Invalid")
        
        try:
            analysis, meta = optimizer.analyze(resume_text, job_desc)
            rewritten = optimizer.rewrite(resume_text, job_desc)
            
            pct = meta.get("percentage", 0)
            status_color = "🟢" if pct >= 70 else "🟡" if pct >= 50 else "🔴"
            
            status = f"{status_color} **Match:** {pct}% | Resume: {len(resume_text):,} chars | Job: {len(job_desc):,} chars"
            
            return analysis, rewritten, status
            
        except Exception as e:
            return (f"## Error\n\n{str(e)}", "Failed", "Status: Error")


# Initialize app
app = ResumeApp()

# =============================================================================
# GRADIO UI
# =============================================================================

def create_interface():
    """Create Gradio interface"""
    
    css = """
    body { font-family: 'Inter', sans-serif; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); }
    .gradio-container { max-width: 1400px !important; }
    .header { text-align: center; padding: 2rem; color: white; }
    .header h1 { font-size: 3rem; font-weight: 800; margin: 0; }
    .main-card { background: white; border-radius: 16px; padding: 2rem; margin-bottom: 1rem; }
    .analyze-btn { background: linear-gradient(135deg, #667eea 0%, #764ba2 100%) !important; color: white !important; 
                   font-weight: 600 !important; padding: 1rem !important; border-radius: 8px !important; 
                   border: none !important; width: 100%; margin-top: 1rem; }
    .status-bar { background: #f1f5f9; padding: 1rem; border-radius: 8px; text-align: center; margin-top: 1rem; }
    .resume-output { background: #fafafa; border: 1px solid #e2e8f0; border-radius: 8px; padding: 1.5rem; 
                     white-space: pre-wrap; font-family: Georgia, serif; line-height: 1.6; }
    """
    
    with gr.Blocks(title="Naked Resume - AI Analyzer") as demo:
        components = {}
        
        gr.HTML('<div class="header"><h1>NAKED RESUME</h1><p>Universal Resume Analyzer</p></div>')
        
        with gr.Row():
            with gr.Column(scale=1):
                with gr.Column(elem_classes="main-card"):
                    gr.Markdown("### 📄 Upload Resume")
                    components['resume'] = gr.File(label="PDF or DOCX", file_types=[".pdf", ".docx"])
                    
                    gr.Markdown("### 🎯 Job Description")
                    components['job'] = gr.Textbox(label="Job Description", lines=10, 
                                                    placeholder="Paste complete job description...")
                    
                    components['btn'] = gr.Button("🔍 Analyze & Optimize", elem_classes="analyze-btn")
                    components['status'] = gr.Markdown(elem_classes="status-bar")
            
            with gr.Column(scale=1):
                with gr.Column(elem_classes="main-card"):
                    gr.Markdown("### 📊 Analysis")
                    components['analysis'] = gr.Markdown()
        
        with gr.Row():
            with gr.Column():
                with gr.Column(elem_classes="main-card"):
                    gr.Markdown("### ✨ Optimized Resume")
                    components['rewrite'] = gr.Markdown(elem_classes="resume-output")
        
        components['btn'].click(
            fn=app.process,
            inputs=[components['resume'], components['job']],
            outputs=[components['analysis'], components['rewrite'], components['status']]
        )
    
    return demo, css


# =============================================================================
# ENTRY POINT
# =============================================================================

if __name__ == "__main__":
    print("=" * 60)
    print("NAKED RESUME - Starting")
    print(f"API Key: {config.OPENAI_API_KEY[:15]}...")
    print(f"Key Type: {'Project Key' if 'proj' in config.OPENAI_API_KEY else 'Standard Key'}")
    print(f"Client Ready: {ai_client.is_ready()}")
    if ai_client.error:
        print(f"Error: {ai_client.error}")
    print("=" * 60)
    
    demo, css_code = create_interface()
    demo.launch(server_name="0.0.0.0", server_port=7860, css=css_code)
